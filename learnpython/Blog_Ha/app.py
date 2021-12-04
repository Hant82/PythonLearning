# Import Flask class và hàm render_template để render file HTML
#from flask import Flask, render_template, request
from flask import Flask, render_template, Response, request, redirect
from pymysql import connect, cursors, Error
from datetime import datetime  
from docx import Document
from docx.shared import Inches
#from app import app

config = {
  'host': 'localhost',
  'user': 'root',
  'password': '1234335',
  'database': 'blog_ha',
  }
cnx = connect(**config)


# Khởi tạo Flask app và kiểm tra xem nó là main script hay imported
app = Flask(__name__)
# Chỉ định URL kích hoạt hàm homepage
@app.route("/")
def get_homepage():
    cur = cnx.cursor()
    sql="SELECT * from bai_viet order by date_time desc"
    cur.execute(sql)
    return render_template("index.html",cur=cur)


@app.route("/", methods=["POST"])
# Hàm homepage() chạy khi đường dẫn khớp với route /
def homepage():
    title = request.form["title"]
    content = request.form["content"] 
    cur = cnx.cursor()
    sql="INSERT INTO bai_viet (ten, noi_dung, date_time) VALUES (%s, %s,%s)"
    value=(title,content,datetime.now())
    try:
        cur.execute(sql,value)
        cnx.commit()
    except:
        cnx.rollback()
    #cnx.close()
    return redirect("/", code=302)


@app.route("/edit/<id>")
def get_edit(id):
    cur = cnx.cursor()
    sql="SELECT * from bai_viet where id="+str(id)
    cur.execute(sql)
    for i in cur:
        r0=i[0]
        r1=i[1]
        r2=i[2]
    return render_template("edit.html",r0=r0,r1=r1,r2=r2)


@app.route("/edit/<id>", methods=["POST"])
# Hàm homepage() chạy khi đường dẫn khớp với route /
def edit(id):
    title = request.form["title"]
    content = request.form["content"] 
    cur = cnx.cursor()
    sql=f"UPDATE bai_viet SET ten='{title}',noi_dung='{content}' where id = {id}"
    # value=(title,content,datetime.now())
    # try:
    cur.execute(sql)
    cnx.commit()
    # except:
    #     cnx.rollback()
    # #cnx.close()
    return redirect("/", code=302)

@app.route("/letter")
def get_letter():
    # cur = cnx.cursor()
    # sql="SELECT * from bai_viet where id="+str(id)
    # cur.execute(sql)
    # for i in cur:
    #     r0=i[0]
    #     r1=i[1]
    #     r2=i[2]
    return render_template("letter.html")


@app.route("/letter", methods=["POST"])
# Hàm homepage() chạy khi đường dẫn khớp với route /
def letter():
    title = request.form["title"]
    content = request.form["content"] 
    document = Document()
    document.add_heading(title, 0)
    document.add_paragraph(content)
    document.save('static/demo.docx')
    return redirect("/letter", code=302)

@app.route("/about")
def get_about():
    # cur = cnx.cursor()
    # sql="SELECT * from bai_viet where id="+str(id)
    # cur.execute(sql)
    # for i in cur:
    #     r0=i[0]
    #     r1=i[1]
    #     r2=i[2]
    return render_template("about.html")


# # Kiểm tra nếu là main script
if __name__ == "__main__":
    app.run(debug=True)
